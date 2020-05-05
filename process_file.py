import os
import docx
import re

# Const STYLE_MULTIPLECHOICEQ = "ВопрМножВыбор"
# Const STYLE_RIGHT_ANSWER = "ВерныйОтвет"
# Const STYLE_WRONG_ANSWER = "НеверныйОтвет"
question_style = "ВопрМножВыбор"
question_numeric_style = "ВопрЧисловой"
right_style = "ВерныйОтвет"
wrong_style = "НеверныйОтвет"
category_style = "Категория"

# regexp
answer_star = re.compile(r"\s*\*+\s*")
empty_pat = re.compile(r"")
numeration_question = re.compile(r"\s*[А-Яа-яA-Za-z]\s")
numeration_answer = re.compile(r"")

# markers
category_marker = "##категория##"


def validate(input_folder, filename):
    input_filename = os.path.join(input_folder, filename)

    print(f"opening {input_filename}")
    try:
        document = docx.Document(input_filename)
        para = document.add_paragraph("TEST")
        para.style = question_style
        para.style = right_style
        para.style = wrong_style
    except Exception as e:
        print(e)
        print("Invalid file")
        return False
    print("OK")
    return True


def paragraph_empty(para):
    # либо 2 run
    # либо 1 run с непечатаемым текстом
    # критерии пустоты:
    # 0 runs
    # 1+ runs && para.text
    return para.text


def process_if_category(para):
    if para.text.lower().startswith(category_marker):
        para.text = para.text.replace(category_marker, "")
        para.style = category_style
        return True
    return False


def process(input_folder, filename, output_folder, numeric=False):
    input_filename = os.path.join(input_folder, filename)
    output_filename = os.path.join(output_folder, filename)

    numeric_answer_pattern = re.compile(r"\s*[Оо]твет\s*[:;]*\s*")

    print(f"opening {input_filename}")
    document = docx.Document(input_filename)

    prev_was_empty = True
    current_empty = False

    paragraphs = document.paragraphs

    if numeric:
        print("Numeric file")
        for para in paragraphs:
            current_empty = len(para.runs) == 0
            if not current_empty:
                if process_if_category(para):
                    current_empty = True
                elif prev_was_empty:  # вопрос
                    para.style = question_numeric_style
                else:  # ответ
                    txt = para.text
                    if txt and re.match(numeric_answer_pattern, txt):
                        para.text = re.sub(numeric_answer_pattern, "", txt, 1)
                        # if len(para.runs[0].text) == 0:
                        #     para.runs[0].text = ""
                        ans = para.text
                        ans = ans.replace(",", ".", 1).strip()
                        para.text = ans
                        para.style = right_style  # правильный
                    else:
                        para.style = wrong_style  # неправильный

            prev_was_empty = current_empty

    else:
        print("Multi choice file")
        # for para in paragraphs:
        #     print(para._p.xml)
        #     continue
        #     if 'a:graphicData' in para._p.xml:
        #         print("img")
        #     else:
        #         print("text")
        # exit()

        for para in paragraphs:
            current_empty = len(para.runs) == 0
            # print(f"Runs: {len(para.runs)} considered {'empty' if current_empty else 'not empty'}")
            # print(f"Text: {' '.join(elem.text for elem in para.runs)}")
            if not current_empty:
                if process_if_category(para):
                    current_empty = True
                elif prev_was_empty:  # вопрос
                    para.style = question_style
                else:               # ответ
                    txt = para.runs[0].text
                    if txt and txt[0] == "*":  # todo: использовать регексп
                        para.runs[0].text = re.sub(answer_star, "", txt, 1)
                        if len(para.runs[0].text) == 0:
                            para.runs[0].text = " "
                        para.style = right_style  # правильный
                    else:
                        para.style = wrong_style  # неправильный

            prev_was_empty = current_empty
    print("Saving document...")
    document.save(output_filename)
    print("Document saved.")
    # TODO: убирать звёздочки

# inp = r"D:\Programming Projects\P\moodleconverter\static\files\input\\"
# out = r"D:\Programming Projects\P\moodleconverter\static\files\output\\"
# # name = r"de6ace26e736446f98e0f4b4a4b1e7a5.docx"
# name = r"ит в химии макро.docx"
#
# process(inp, name, out)
# validate(inp, name)